import zipfile
import tempfile
from pathlib import Path
import shutil
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import py7zr
import rarfile

def create_test_files():
    test_dir = Path("test_data")
    test_dir.mkdir(exist_ok=True)
    
    print("Creating test files with Pushkin texts in test_data/")
    
    pushkin_texts = {
        "hello.txt": "Мороз и солнце; день чудесный!\nЕще ты дремлешь, друг прелестный —\nПора, красавица, проснись...",
        "latin.txt": "I remember a wonderful moment:\nYou appeared before me,\nLike a fleeting vision,\nLike a genius of pure beauty.",
        "note.txt": "У лукоморья дуб зелёный;\nЗлатая цепь на дубе том:\nИ днём и ночью кот учёный\nВсё ходит по цепи кругом.",
        "deep.txt": "Я помню чудное мгновенье:\nПередо мной явилась ты,\nКак мимолётное виденье,\nКак гений чистой красоты.",
        "text_in_zip.txt": "Сказка о царе Салтане:\nТри девицы под окном\nПряли поздно вечерком.",
        "readme.txt": "Евгений Онегин:\nМой дядя самых честных правил,\nКогда не в шутку занемог,\nОн уважать себя заставил\nИ лучше выдумать не мог."
    }
    
    for filename, content in pushkin_texts.items():
        with open(test_dir / filename, "w", encoding='utf-8') as f:
            f.write(content)
    
    doc = Document()
    doc.add_heading('Александр Сергеевич Пушкин', 0)
    doc.add_paragraph('Я помню чудное мгновенье:')
    doc.add_paragraph('Передо мной явилась ты,')
    doc.add_paragraph('Как мимолётное виденье,')
    doc.add_paragraph('Как гений чистой красоты.')
    doc.add_paragraph('')
    doc.add_paragraph('В томленьях грусти безнадежной,')
    doc.add_paragraph('В тревогах шумной суеты,')
    doc.add_paragraph('Звучал мне долго голос нежный')
    doc.add_paragraph('И снились милые черты.')
    doc.save(test_dir / "pushkin.docx")
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Герои Пушкина"
    data = [
        ["Персонаж", "Произведение", "Год"],
        ["Евгений Онегин", "Евгений Онегин", 1823],
        ["Татьяна Ларина", "Евгений Онегин", 1823],
        ["Пётр Гринёв", "Капитанская дочка", 1836],
        ["Пугачёв", "Капитанская дочка", 1836],
        ["Германн", "Пиковая дама", 1834],
        ["Дубровский", "Дубровский", 1833],
        ["Борис Годунов", "Борис Годунов", 1825],
        ["Руслан", "Руслан и Людмила", 1820],
        ["Людмила", "Руслан и Людмила", 1820],
    ]
    
    for row in data:
        ws.append(row)
    
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column].width = adjusted_width
    
    wb.save(test_dir / "pushkin_heroes.xlsx")
    
    c = canvas.Canvas(str(test_dir / "ruslan.pdf"), pagesize=A4)
    c.setFont('Helvetica', 16)
    c.drawString(100, 800, "Ruslan and Ludmila")
    c.setFont('Helvetica', 12)
    
    text_lines = [
        "By the seashore a green oak towers,",
        "With a golden chain upon it bound:",
        "And day and night a learned cat",
        "Paces on that chain round and round;",
        "To the right he goes - he sings a song,",
        "To the left he goes - he tells a tale.",
        "",
        "There wonders happen: a wood-goblin strays,",
        "A mermaid sits among the branches;"
    ]
    
    y = 760
    for line in text_lines:
        c.drawString(100, y, line)
        y -= 20
    
    c.save()
    print("ruslan.pdf created (English version)")
    
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        
        for filename, content in pushkin_texts.items():
            with open(tmp_path / filename, "w", encoding='utf-8') as f:
                f.write(content)
        
        shutil.copy(test_dir / "pushkin.docx", tmp_path / "pushkin.docx")
        shutil.copy(test_dir / "pushkin_heroes.xlsx", tmp_path / "heroes.xlsx")
        shutil.copy(test_dir / "ruslan.pdf", tmp_path / "ruslan.pdf")
        
        print("  Creating ZIP archive with Pushkin texts...")
        with zipfile.ZipFile(test_dir / "pushkin.zip", "w") as zipf:
            for file in tmp_path.iterdir():
                zipf.write(file, file.name)
        
        print("  Creating 7z archive with Pushkin texts...")
        with py7zr.SevenZipFile(test_dir / "pushkin.7z", "w") as sz:
            for file in tmp_path.iterdir():
                sz.write(file, file.name)
        
        print("  Creating RAR archive with Pushkin texts...")
        try:
            with rarfile.RarFile(test_dir / "pushkin.rar", "w") as rf:
                for file in tmp_path.iterdir():
                    rf.write(file, file.name)
        except Exception as e:
            print(f"  RAR not created: {e}")
    
    print("  Creating nested ZIP...")
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        
        inner_zip = tmp_path / "pushkin_inside.zip"
        with zipfile.ZipFile(inner_zip, "w") as inner:
            inner.writestr("onegin.txt", "Евгений Онегин:\nМой дядя самых честных правил...")
        
        with zipfile.ZipFile(test_dir / "pushkin_nested.zip", "w") as outer:
            outer.write(inner_zip, "pushkin_inside.zip")
    
    poems_dir = test_dir / "poems"
    poems_dir.mkdir(exist_ok=True)
    
    poems = {
        "winter_morning.txt": "Зимнее утро:\nМороз и солнце; день чудесный!",
        "to_kerne.txt": "К***:\nЯ помню чудное мгновенье...",
        "prophet.txt": "Пророк:\nДуховной жаждою томим...",
    }
    
    for filename, content in poems.items():
        with open(poems_dir / filename, "w", encoding='utf-8') as f:
            f.write(content)
    
    print("\nTest files with Pushkin texts created successfully!")
    print(f"Location: {test_dir.absolute()}")
    
    print("\nCreated files:")
    for f in sorted(test_dir.glob("*")):
        if f.is_file():
            size = f.stat().st_size
            print(f"  {f.name} ({size} bytes)")
        elif f.is_dir():
            print(f"  {f.name}/")
            for subf in sorted(f.glob("*")):
                subsize = subf.stat().st_size
                print(f"      {subf.name} ({subsize} bytes)")

if __name__ == "__main__":
    create_test_files()